import streamlit as st
import pandas as pd
import re
import os
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ==========================================
# CONFIGURAÇÃO DA PÁGINA STREAMLIT
# ==========================================
st.set_page_config(page_title="Gerador de Relatórios DP", layout="wide", page_icon="⚓")

# Estilo CSS para manter o padrão visual azul
st.markdown("""
    <style>
    :root { --primary-color: #1F4E79; }
    div.stButton > button:first-child {
        background-color: #1F4E79;
        color: white;
        border-radius: 5px;
        border: none;
        font-weight: bold;
    }
    div.stButton > button:hover { background-color: #163a5c; color: white; }
    .stSidebar { background-color: #f0f2f6; }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# BARRA LATERAL (CONFIGURAÇÕES)
# ==========================================
st.sidebar.header("⚙️ Configurações do Documento")

# 1. Configuração da Planilha
st.sidebar.subheader("Dados de Entrada")
sheet_input = st.sidebar.text_input(
    "Aba da Planilha (Nome ou Índice)", 
    value="0", 
    help="Digite '0' para a primeira aba, '1' para a segunda, ou o nome exato da aba (ex: 'Dados'). Se o código der erro, tente mudar esse número."
)

# 2. Configuração de Estilo
st.sidebar.subheader("Estilo e Fontes")
font_name_cfg = st.sidebar.selectbox(
    "Fonte Principal", 
    ["Raleway", "Arial", "Calibri", "Times New Roman"], 
    index=0,
    help="Define a fonte de todo o documento. 'Raleway' é o padrão do design original."
)

h1_size_cfg = st.sidebar.number_input("Tamanho Título Principal (H1)", value=20, step=1)
h2_size_cfg = st.sidebar.number_input("Tamanho Título Seção (H2)", value=16, step=1)
body_size_cfg = st.sidebar.number_input("Tamanho Texto Corpo", value=10, step=1, help="Tamanho padrão do texto dentro das tabelas e descrições.")
table_font_size_cfg = st.sidebar.number_input("Tamanho Texto Tabela (Pequeno)", value=9, step=1, help="Usado em rótulos e cabeçalhos de tabelas densas.")

# 3. Margens
st.sidebar.subheader("Margens da Página")
margin_top = st.sidebar.slider("Margem Superior (pol)", 0.5, 2.0, 1.0, 0.1)
margin_bottom = st.sidebar.slider("Margem Inferior (pol)", 0.5, 2.0, 0.8, 0.1)
margin_side = st.sidebar.slider("Margens Laterais (pol)", 0.2, 1.5, 0.5, 0.1)

# Agrupando configurações para passar para as funções
STYLE_CONFIG = {
    "font_name": font_name_cfg,
    "h1": h1_size_cfg,
    "h2": h2_size_cfg,
    "body": body_size_cfg,
    "small": table_font_size_cfg,
    "margins": (margin_top, margin_bottom, margin_side, margin_side)
}

# ==========================================
# CONSTANTES GERAIS
# ==========================================
COLOR_PRIMARY = "1F4E79"
COLOR_BORDER  = "BFBFBF"
COLOR_BG_UNIFIED = "F2F2F2"
COLOR_TEXT_MAIN = RGBColor(0x26, 0x26, 0x26)
COLOR_TEXT_LABEL = RGBColor(0x1F, 0x4E, 0x79)
COLOR_TEXT_PLACEHOLDER = RGBColor(89, 89, 89)
LOGO_PATH = 'logo.png' 

refined_border = {"sz": 8, "val": "single", "color": COLOR_BORDER}
box_border_settings = {
    "top": refined_border, "bottom": refined_border, "left": refined_border, "right": refined_border
}
no_border = {
    "top": {"sz": 0, "val": "nil", "color": "auto"},
    "bottom": {"sz": 0, "val": "nil", "color": "auto"},
    "left": {"sz": 0, "val": "nil", "color": "auto"},
    "right": {"sz": 0, "val": "nil", "color": "auto"},
    "insideV": {"sz": 0, "val": "nil", "color": "auto"}
}

# ==========================================
# FUNÇÕES WORD
# ==========================================

def set_cell_border_and_shading(cell, border_settings=None, shading_color=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for element in tcPr.xpath('./w:tcBorders'): tcPr.remove(element)

    if border_settings:
        tcBorders = OxmlElement('w:tcBorders')
        for edge, data in border_settings.items():
            element = OxmlElement(f"w:{edge}")
            for key, value in data.items(): element.set(qn(f"w:{key}"), str(value))
            tcBorders.append(element)
        tcPr.append(tcBorders)

    for element in tcPr.xpath('./w:shd'): tcPr.remove(element)
    if shading_color:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), shading_color)
        tcPr.append(shd)

def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, value in kwargs.items():
        mar = OxmlElement(f'w:{edge}')
        mar.set(qn('w:w'), str(value)); mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)

def set_table_indent(table, indent_val=0):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    for el in tblPr.xpath("w:tblLayout"): tblPr.remove(el)
    tblPr.append(layout)
    for element in tblPr.xpath('./w:tblCellSpacing'): tblPr.remove(element)
    tblCellSpacing = OxmlElement('w:tblCellSpacing'); tblCellSpacing.set(qn('w:w'), "0"); tblCellSpacing.set(qn('w:type'), "dxa")
    tblPr.append(tblCellSpacing)
    for el in tblPr.xpath("w:tblInd"): tblPr.remove(el)
    tblInd = OxmlElement('w:tblInd'); tblInd.set(qn('w:w'), str(indent_val)); tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

def create_header(doc, image_path):
    section = doc.sections[0]
    section.header_distance = Inches(0.2)
    header = section.header
    for paragraph in header.paragraphs:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    p_logo = header.add_paragraph()
    p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_logo.paragraph_format.left_indent = Inches(-0.09)
    if os.path.exists(image_path):
        run_logo = p_logo.add_run()
        run_logo.add_picture(image_path, width=Inches(7.5))
    else:
        run_logo = p_logo.add_run("[LOGO NÃO ENCONTRADO - Verifique repositório]")
        run_logo.font.color.rgb = RGBColor(255, 0, 0); run_logo.font.size = Pt(8)

def create_details_section(doc, test_info, cfg):
    details_table = doc.add_table(rows=0, cols=1)
    details_table.width = Inches(7.5)
    details_table.allow_autofit = False
    set_table_indent(details_table, indent_val=0)

    def _add_box_row(label, content, shading_color, is_placeholder=False):
        row = details_table.add_row()
        cell = row.cells[0]; cell.width = Inches(7.5)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=100)
        set_cell_border_and_shading(cell, border_settings=box_border_settings, shading_color=shading_color)
        
        if label:
            p_label = cell.paragraphs[0]
            if not p_label.text: p_label.clear()
            p_label.paragraph_format.space_after = Pt(3)
            run_label = p_label.add_run(f"{label}:"); run_label.bold = True
            run_label.font.name = cfg['font_name']
            run_label.font.size = Pt(cfg['small']); run_label.font.color.rgb = COLOR_TEXT_LABEL
        
        p_content = cell.add_paragraph() if label else cell.paragraphs[0]
        p_content.paragraph_format.space_after = Pt(2)
        
        if is_placeholder:
            run_content = p_content.add_run(content); run_content.font.italic = True
            run_content.font.color.rgb = COLOR_TEXT_PLACEHOLDER; run_content.font.size = Pt(cfg['small'])
        else:
            lines = str(content).split('\n')
            for i, line in enumerate(lines):
                line = line.strip()
                if not line: continue
                match = re.match(r'^(\d+\.)\s*(.*)', line)
                if match:
                    r_num = p_content.add_run(match.group(1) + " "); r_num.font.size = Pt(cfg['body']); r_num.bold = True
                    r_txt = p_content.add_run(match.group(2)); r_txt.font.size = Pt(cfg['body']); r_txt.bold = False
                else:
                    r_line = p_content.add_run(line); r_line.font.size = Pt(cfg['body'])
                
                # Aplica fonte configurada
                for run in p_content.runs: run.font.name = cfg['font_name']; run.font.color.rgb = COLOR_TEXT_MAIN
                
                if i < len(lines) - 1: p_content.add_run("\n")
        return p_content

    if test_info.get('Objective'): _add_box_row("Objective", str(test_info['Objective']), COLOR_BG_UNIFIED)
    _add_box_row("Method", test_info['Method'], COLOR_BG_UNIFIED)
    _add_box_row("Steps", "\n".join(str(step) for step in test_info['Steps']), COLOR_BG_UNIFIED)
    _add_box_row("Expected Results", "\n".join(map(str, test_info['Expected Results'])), COLOR_BG_UNIFIED)
    
    results_content = "\n".join(str(r).strip() for r in test_info.get('Result + Comment', []) if pd.notna(r) and str(r).strip().lower() != "nan")
    is_ph = False if results_content else True
    if not results_content: results_content = "No results or comments provided."
    _add_box_row("Results", results_content, COLOR_BG_UNIFIED, is_placeholder=is_ph)

    comments_list = test_info.get('Step Comments', [])
    row = details_table.add_row()
    cell = row.cells[0]; cell.width = Inches(7.5)
    set_cell_margins(cell, top=80, bottom=80, left=120, right=100)
    set_cell_border_and_shading(cell, border_settings=box_border_settings, shading_color=COLOR_BG_UNIFIED)
    
    p_label = cell.paragraphs[0]
    if not p_label.text: p_label.clear()
    p_label.paragraph_format.space_after = Pt(3)
    run_label = p_label.add_run("Comments:"); run_label.bold = True; run_label.font.size = Pt(cfg['small']); run_label.font.color.rgb = COLOR_TEXT_LABEL
    run_label.font.name = cfg['font_name']

    p_content = cell.add_paragraph()
    if not comments_list:
        run_ph = p_content.add_run("No additional comments"); run_ph.font.italic = True; run_ph.font.size = Pt(cfg['small'])
        run_ph.font.color.rgb = COLOR_TEXT_PLACEHOLDER; run_ph.font.name = cfg['font_name']
    else:
        for i, item in enumerate(comments_list):
            run_step = p_content.add_run(f"Step {item['step']}: "); run_step.bold = True; run_step.font.size = Pt(cfg['body'])
            run_text = p_content.add_run(f"{item['text']}"); run_text.bold = False; run_text.font.size = Pt(cfg['body'])
            for r in [run_step, run_text]: r.font.name = cfg['font_name']; r.font.color.rgb = COLOR_TEXT_MAIN
            if i < len(comments_list) - 1: p_content.add_run("\n")

def create_test_page(doc, test_info, cfg, is_first_test=False, section_title=None):
    if not is_first_test and len(doc.paragraphs) > 0: doc.add_page_break()

    if section_title:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(6); p_sec.paragraph_format.space_after = Pt(14)
        run_sec = p_sec.add_run(str(section_title))
        run_sec.font.name = cfg['font_name']; run_sec.font.size = Pt(cfg['h2']); run_sec.font.color.rgb = RGBColor(0,0,0); run_sec.bold = True

    # Header Tabela Azul
    table_blue = doc.add_table(rows=1, cols=2); table_blue.width = Inches(7.5); table_blue.allow_autofit = False
    set_table_indent(table_blue, indent_val=-10)
    table_blue.columns[0].width = Inches(0.8); table_blue.columns[1].width = Inches(6.7)
    
    cell_lbl = table_blue.cell(0, 0); cell_val = table_blue.cell(0, 1)
    set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
    set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

    p_lbl = cell_lbl.paragraphs[0]; run_lbl = p_lbl.add_run("TEST NO:")
    run_lbl.bold = True; run_lbl.font.size = Pt(cfg['small'] + 1); run_lbl.font.color.rgb = RGBColor(255, 255, 255)
    
    p_val = cell_val.paragraphs[0]; run_val = p_val.add_run(f"{test_info['Test']}")
    run_val.bold = True; run_val.font.size = Pt(cfg['small'] + 1); run_val.font.color.rgb = RGBColor(255, 255, 255)
    
    for r in [run_lbl, run_val]: r.font.name = cfg['font_name']

    set_cell_border_and_shading(cell_lbl, border_settings=no_border, shading_color=COLOR_PRIMARY)
    set_cell_border_and_shading(cell_val, border_settings=no_border, shading_color=COLOR_PRIMARY)
    
    # Gap minúsculo (1pt) apenas para separar as tabelas sem mesclar
    p_gap_1 = doc.add_paragraph()
    p_gap_1.paragraph_format.space_before = Pt(0)
    p_gap_1.paragraph_format.space_after = Pt(0)
    p_gap_1.paragraph_format.line_spacing = Pt(1)
    p_gap_1.add_run().font.size = Pt(1)

    # Tabela FMEA
    table_info = doc.add_table(rows=1, cols=2); table_info.width = Inches(7.49); set_table_indent(table_info, indent_val=0)
    table_info.columns[0].width = Inches(3.75); table_info.columns[1].width = Inches(3.74)
    
    c1 = table_info.cell(0,0); c2 = table_info.cell(0,1)
    set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
    set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
    
    p1 = c1.paragraphs[0]; r1a = p1.add_run("FMEA Reference: "); r1a.bold = True
    r1b = p1.add_run(str(test_info.get('FMEA Reference', '-')))
    
    p2 = c2.paragraphs[0]; r2a = p2.add_run("Sub-System: "); r2a.bold = True
    r2b = p2.add_run(str(test_info.get('Sub-System', '-')))

    for r in [r1a, r1b, r2a, r2b]: 
        r.font.size = Pt(cfg['small']); r.font.color.rgb = COLOR_TEXT_MAIN; r.font.name = cfg['font_name']

    set_cell_border_and_shading(c1, border_settings=box_border_settings, shading_color=COLOR_BG_UNIFIED)
    set_cell_border_and_shading(c2, border_settings=box_border_settings, shading_color=COLOR_BG_UNIFIED)

    create_details_section(doc, test_info, cfg)
    
    # --- CORREÇÃO DO ESPAÇO INDESEJADO AQUI ---
    # Substituímos o parágrafo de espaçamento padrão por um "ghost paragraph" de 1pt.
    # Isso evita que o Word funda as tabelas, mas remove o espaço visual branco.
    p_gap_footer = doc.add_paragraph()
    p_gap_footer.paragraph_format.space_before = Pt(0)
    p_gap_footer.paragraph_format.space_after = Pt(0)
    p_gap_footer.paragraph_format.line_spacing = Pt(1) # Linha minúscula
    p_gap_footer.add_run().font.size = Pt(1) # Fonte minúscula

    # Assinaturas
    table_wit = doc.add_table(rows=1, cols=2); table_wit.width = Inches(7.5); set_table_indent(table_wit, indent_val=-10)
    table_wit.columns[0].width = Inches(5.0); table_wit.columns[1].width = Inches(2.5)
    
    cw = table_wit.cell(0,0); cd = table_wit.cell(0,1)
    set_cell_margins(cw, top=60, bottom=60, left=100, right=100)
    set_cell_margins(cd, top=60, bottom=60, left=100, right=100)
    
    pw = cw.paragraphs[0]; rw1 = pw.add_run("Witnessed by: "); rw1.bold = True
    rw2 = pw.add_run(str(test_info.get('Witness 1', '-'))); rw2.bold = True
    
    pd_ = cd.paragraphs[0]; rd1 = pd_.add_run("Date: "); rd1.bold = True
    raw_d = test_info.get('Date:'); date_v = str(raw_d).strip() if pd.notna(raw_d) and str(raw_d).lower()!='nan' else '-'
    rd2 = pd_.add_run(date_v); rd2.bold = True
    
    for r in [rw1, rw2, rd1, rd2]:
        r.font.size = Pt(cfg['small']); r.font.color.rgb = RGBColor(255, 255, 255); r.font.name = cfg['font_name']
        
    set_cell_border_and_shading(cw, border_settings=no_border, shading_color=COLOR_PRIMARY)
    set_cell_border_and_shading(cd, border_settings=no_border, shading_color=COLOR_PRIMARY)

def generate_professional_docx(uploaded_file, cfg):
    # Lógica de seleção de aba com tratamento de erro
    sheet_val = cfg['sheet_target']
    try:
        # Tenta converter para inteiro (índice)
        sheet_target = int(sheet_val)
    except ValueError:
        # Se falhar, usa como string (nome da aba)
        sheet_target = sheet_val

    try:
        df = pd.read_excel(uploaded_file, sheet_name=sheet_target, header=0, dtype={'test number': str})
    except ValueError:
        st.error(f"❌ Erro: Não foi possível encontrar a aba '{sheet_target}'. Verifique se o nome/índice está correto na barra lateral.")
        return None
    except Exception as e:
        st.error(f"❌ Erro ao ler o arquivo Excel: {e}")
        return None

    # Verifica colunas mínimas
    required_cols = ['test number', 'Section', 'Test', 'Method', 'Step', 'Expected Result']
    missing = [c for c in required_cols if c not in df.columns]
    if missing:
        st.error(f"❌ A aba selecionada não possui as colunas obrigatórias: {missing}. Tente outra aba.")
        return None

    grouped_tests = {}
    if not df.empty:
        for _, row in df.iterrows():
            chapter_title = row.get('Section', 'Section Name')
            test_number = row.get('test number', '000')
            if pd.isna(test_number): continue

            if test_number not in grouped_tests:
                grouped_tests[test_number] = {
                    'Test': row.get('Test', 'Test Title'), 'Method': row.get('Method', ''), 'Steps': [],
                    'Expected Results': [], 'Result + Comment': [], 'Step Comments': [],
                    'Witness 1': row.get('Witness 1', ''), 'Date:': row.get('Date', ''), 'Section': chapter_title,
                    'FMEA Reference': row.get('FMEA Reference', ''), 'Sub-System': row.get('Sub-System', ''),
                    'Objective': row.get('Objective', '')
                }
            grouped_tests[test_number]['Steps'].append(row.get('Step', ''))
            
            # Comentários
            curr_step = len(grouped_tests[test_number]['Steps'])
            a_comm = row.get('Auditor FMEA Comment')
            if pd.notna(a_comm) and str(a_comm).strip():
                grouped_tests[test_number]['Step Comments'].append({'step': curr_step, 'text': str(a_comm).strip()})
            
            grouped_tests[test_number]['Expected Results'].append(row.get('Expected Result', ''))
            grouped_tests[test_number]['Result + Comment'].append(row.get('Result + Comment', ''))

    # Configuração Inicial do Doc
    doc = Document()
    normal_style = doc.styles['Normal']
    normal_style.font.name = cfg['font_name']
    normal_style.font.size = Pt(cfg['body'])
    # --- CORREÇÃO CRÍTICA: ZERAR ESPAÇAMENTO PADRÃO ---
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0) # Evita gaps automáticos do Word
    normal_style.paragraph_format.line_spacing = 1.15

    section = doc.sections[0]
    section.top_margin = Inches(cfg['margins'][0])
    section.bottom_margin = Inches(cfg['margins'][1])
    section.left_margin = Inches(cfg['margins'][2])
    section.right_margin = Inches(cfg['margins'][3])
    section.page_width = Inches(8.5); section.page_height = Inches(11.0)

    create_header(doc, LOGO_PATH)

    main_title_para = doc.add_paragraph()
    main_title_run = main_title_para.add_run("APPENDIX B - DP ANNUAL TRIALS TESTS")
    main_title_run.font.name = cfg['font_name']
    main_title_run.font.size = Pt(cfg['h1'])
    main_title_run.font.color.rgb = RGBColor(0, 0, 0); main_title_run.bold = True
    main_title_para.paragraph_format.space_after = Pt(24)

    first_iteration = True
    current_chapter = None
    
    if grouped_tests:
        for test_number, test_info in grouped_tests.items():
            test_info['test number'] = test_number
            if test_info['Section'] != current_chapter:
                current_chapter = test_info['Section']
                create_test_page(doc, test_info, cfg, is_first_test=first_iteration, section_title=current_chapter)
            else:
                create_test_page(doc, test_info, cfg, is_first_test=first_iteration, section_title=None)
            first_iteration = False

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# INTERFACE PRINCIPAL
# ==========================================

# Logo na Interface Streamlit
if os.path.exists('bram_logo.png'):
     col1, col2, col3 = st.columns([1, 2, 1])
     with col2: st.image('bram_logo.png', use_container_width=True)
else:
     if os.path.exists(LOGO_PATH):
         col1, col2, col3 = st.columns([1, 2, 1])
         with col2: st.image(LOGO_PATH, use_container_width=True)

st.title("Gerador de Relatórios DP - Padrão Profissional")
st.markdown("Faça o upload da planilha Excel para gerar o relatório formatado (Apêndice B).")
st.info("ℹ️ Utilize a barra lateral (👈) para configurar qual aba do Excel ler, alterar fontes, tamanhos e margens.")

uploaded_file = st.file_uploader("Upload Planilha (.xlsx)", type=["xlsx"])

if uploaded_file:
    # Passamos o dicionário STYLE_CONFIG (que contém os valores do sidebar) para a função
    STYLE_CONFIG['sheet_target'] = sheet_input # Adiciona o input da aba ao config
    
    if st.button("Gerar Relatório DOCX", type="primary"):
        with st.spinner("Lendo planilha e formatando documento..."):
            docx_buffer = generate_professional_docx(uploaded_file, STYLE_CONFIG)
            
            if docx_buffer:
                st.success(f"Relatório gerado com sucesso usando a aba: '{sheet_input}'")
                st.download_button(
                    label="⬇️ Baixar test_report_professional.docx",
                    data=docx_buffer,
                    file_name="test_report_professional.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
